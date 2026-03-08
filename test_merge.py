import docx
import pypandoc

def test_insertion():
    # 1. Generate pandoc docx
    markdown = "# Test Heading\n\nThis is a test paragraph from **Pandoc**.\n\n- List item 1\n- List item 2"
    with open("test.md", "w") as f:
        f.write(markdown)
    pypandoc.convert_file("test.md", "docx", outputfile="pandoc_out.docx")

    # 2. Open both documents
    doc_main = docx.Document("wbtemplate.docx")
    doc_pandoc = docx.Document("pandoc_out.docx")

    # 3. Find markers in main doc
    start_marker_para = None
    end_marker_table = None

    for p in doc_main.paragraphs:
        if "titel eerste hoofdstuk" in p.text.lower():
            start_marker_para = p
            break
            
    for t in doc_main.tables:
        table_text = "".join(cell.text for row in t.rows for cell in row.cells).lower()
        if "deze tekst laten staan" in table_text and "laatste pagina berekening" in table_text:
            end_marker_table = t
            break

    body_main = doc_main.element.body
    body_pandoc = doc_pandoc.element.body

    start_idx = -1
    end_idx = -1
    for i, child in enumerate(body_main):
        if start_marker_para and child is start_marker_para._element:
            start_idx = i
        if end_marker_table and child is end_marker_table._element:
            end_idx = i

    if start_idx != -1 and end_idx != -1:
        tail_elements = []
        for i in range(end_idx, len(body_main)):
            tail_elements.append(body_main[i])
        
        for _ in range(len(body_main) - (start_idx + 1)):
            body_main.remove(body_main[start_idx + 1])

        # Insert pandoc elements (except the sectPr at the very end of body_pandoc!)
        for elem in body_pandoc:
            if elem.tag.endswith('sectPr'):
                continue
            body_main.append(elem)

        # Append tail elements
        for elem in tail_elements:
            body_main.append(elem)

    doc_main.save("merged_test_output.docx")

test_insertion()
