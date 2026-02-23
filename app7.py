import streamlit as st
import io
import os
import tempfile 
from datetime import datetime
import pandas as pd
import re
import json
from collections import Counter
import concurrent.futures
import sys

# Third-party imports
from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader 
from langchain_text_splitters import RecursiveCharacterTextSplitter 
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain_core.runnables import RunnablePassthrough, RunnableParallel
from langchain_core.output_parsers import StrOutputParser
from thefuzz import process
from thefuzz import fuzz
from openai import OpenAI

import markdown
from htmldocx import HtmlToDocx
from docx.shared import Inches

# Utility imports
from utils import RUIS_WOORDEN, generate_csv_from_municipality, get_geodata_for_municipality, create_map_image, PAD_GEMEENTEN


# --- 1. PAGE CONFIGURATION (MUST BE FIRST) ---
st.set_page_config(page_title="Passende beoordeling voor omgevingsvisies", layout="wide")

# --- 2. CONFIGURATION & SECRETS ---
VECTOR_STORE_DIRECTORY = "vector_store"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"

try:
    if "openai_base_url" not in st.session_state:
        st.session_state.openai_base_url = st.secrets["BASE_URL"]
    if "openai_api_key" not in st.session_state:
        st.session_state.openai_api_key = st.secrets["API_KEY"]
        
    YOUR_API_BASE_URL = st.secrets["BASE_URL"]
    YOUR_API_KEY = st.secrets["API_KEY"]
except KeyError:
    st.error("API keys (API_KEY or BASE_URL) not found in `.streamlit/secrets.toml`.")
    st.stop()

# --- 3. PROMPT TEMPLATES ---
SYSTEM_TEMPLATE = """
Je bent een behulpzame, deskundige assistent voor het analyseren van documenten over natuurdoelstellingen. Specifiek gaat dit over instandhoudingsdoelstellingen. Het concept 'instandhoudingsdoelstellingen' staat centraal.
Je taak is om de geleverde context strikt te analyseren.

**STAP 1: CONCEPT CHECK**
Je stelt de aanwezigheid van instandhoudingsdoelstellingen vast op 4 verschillende typen natuur.

**STAP 2: GESTRUCTUREERDE ANALYSE (JSON)**
Je genereert een gedetailleerde analyse in JSON-formaat. Je mag uitsluitend de resultaten gebruiken van de concept-check om te bepalen welke typen natuur je moet analyseren.

Gebruik geen informatie die niet in de geleverde context staat. Wees beknopt en gezaghebbend.

Context:
{context}
"""

CONCEPT_CHECK_PROMPT = """
Analyseer de geleverde context en geef aan of er in de bron instandhoudingsdoelstellingen zijn voor de volgende 4 natuurtypen. Status 'Ja' betekent dat er 1 of meer instandhoudingsdoelstellingen zijn voor dit natuurtype.

- Habitattype: [Status]
- Habitatrichtlijnsoorten: [Status]
- Broedvogels: [Status]
- Niet-broedvogels: [Status]

Geef geen andere tekst.
"""

TABLE_GENERATION_PROMPT = """
Jouw taak is om een gedetailleerde analyse te genereren over de natuurdoelen.
**Geef je antwoord UITSLUITEND als een valide JSON-object.** Gebruik geen Markdown-opmaak (zoals ```json) rondom het object.

**GEBRUIK DE VOLGENDE INSTRUCTIES OM TE BEPALEN WELKE CATEGORIEN JE MOET ANALYSEREN. ANALYSEER ALLEEN DE CATEGORIEËN DIE IN DE INSTRUCTIES ALS 'AANWEZIG' ZIJN GEMARKEERD.**
-----------------
INSTRUCTIES/CHECKLIST:
{concept_check_result}
-----------------

Voor elke 'Aanwezige' categorie, en voor elk individueel type binnen die categorie (bijv. elk Habittattype zoals H3150 of elke Broedvogelsoort), voeg je een object toe aan de lijst "bevindingen".

Het JSON-object moet de volgende structuur hebben:
{{
  "bevindingen": [
    {{
      "categorie": "Kies uit: 'Habitattype', 'Habitatrichtlijnsoorten', 'Broedvogels', of 'Niet-broedvogels'",
      "natuurtype": "Bijv: 'H3150 Natuurlijke eutrofe meren' (volledige naam incl code)",
      "kwaliteit": "Korte beschrijving kwaliteit",
      "knelpunten": "Belangrijkste knelpunten",
      "oordeel": "Kies strikt uit: 'Ja', 'Nee, niet haalbaar', of 'Nee, gebrek aan gegevens'"
    }}
  ],
  "samenvatting": "Een algemene samenvatting van maximaal 5 zinnen over de meest opvallende bevindingen."
}}

Als er geen gegevens zijn, laat de lijst "bevindingen" leeg.
"""

# Prompt voor Stap 2 (Omgevingsvisie)
IMPACT_PROMPT_FULL = """
Je bent een expert in ruimtelijke ordening en ecologie.
Hieronder volgt de volledige tekst van een Omgevingsvisie (of beleidsdocument).

**JOUW TAAK:**
Analyseer het document grondig op concrete ingrepen, ambities of ontwikkelingen binnen de volgende 5 categorieën die impact kunnen hebben op de natuur:

1. **Woningbouw:** (Implicaties: ruimtegebruik, extra mensen, extra verkeer, extra recreatiedruk)
2. **Recreatie Ontwikkeling:** (Implicaties: extra recreatiedruk, verstoring)
3. **Mobiliteit & Infrastructuur:** (Implicaties: ruimtegebruik, extra verkeer, versnippering)
4. **Landbouwmaatregelen:** (Implicaties: verplaatsing/intensivering/extensivering kan leiden tot verschuivingen in stikstofdepositie)
5. **Bedrijvigheid:** (Implicaties: extra verkeersbewegingen, ruimtegebruik nieuwe terreinen)

**OUTPUT FORMAAT & RESTRICTIES:**
Genereer voor ELK van deze 5 categorieën een samenvattende tekst.
- Gebruik de categorie als tussenkop (bijv. "#### 1. Woningbouw").
- Beschrijf concreet wat er in het plan staat (aantallen, locaties, specifieke projecten).
- Benoem expliciet de potentiële risico's voor de natuur zoals hierboven beschreven.
- Citeer waar mogelijk paginanummers of paragraafnamen.
- Maak gebruik van bullet points voor de opsommingen binnen een categorie.

**BELANGRIJK:** Sluit de analyse direct af na de 5e categorie. Geef GEEN suggesties voor verdere analyses, stel GEEN wedervragen en bied GEEN extra hulp aan. De output wordt gebruikt in een statisch rapport zonder mogelijkheid tot interactie.

Als er voor een categorie GEEN maatregelen worden genoemd, geef dit dan expliciet aan met "Geen relevante ingrepen gevonden in dit document."

TEKST VAN DOCUMENT:
{context}
"""

# NIEUW: Prompt voor Stap 3 (Conclusies - alleen op basis van impact)
CONCLUSION_PROMPT_TEMPLATE = """
Je bent een expert in ecologie en ruimtelijke ordening.
Je taak is om een concluderende paragraaf te schrijven over een specifiek onderwerp, gebaseerd op de impactanalyse van een omgevingsvisie.

**ONDERWERP:** {topic}

**ANALYSE IMPACT OMGEVINGSVISIE**
Hieronder staat de analyse van de ingrepen uit de omgevingsvisie.
---
{impact_analyse}
---

**JOUW OPDRACHT:**
Schrijf een beknopte, concluderende paragraaf over het opgegeven **ONDERWERP**. 
Focus op de relatie tussen de vijf impact categorieën en de mogelijke effecten op **ONDERWERP**.

Schrijf 1 paragraaf over hoe de impacts uit de vijf categoriëen kunnen leiden tot een verslechtering in de staat van **ONDERWERP**. 
Bijvoorbeeld, een verslechtering in onderwerp 'stikstofdepositie' betekent een toename van de stikstofdepositie, 
een verslechtering in het onderwerp 'verstoring' betekent een toename van verstoring. Onderbouw wel altijd, met informatie uit de impactanalyse, waarom
je deze conclusie maakt. 

Schrijf ook 1 paragraaf over hoe de impacts uit de vijf categoriëen kunnen leiden tot een verbetering in de staat van **ONDERWERP**.
Bijvoorbeeld, een verbetering in het onderwerp 'stikstofdepositie' betekent een afname van de stikstofdepositie, of
een verbetering in het onderwerp 'verstoring' betekent een afname van de verstoring. Onderbouw wel altijd, met informatie uit de impactanalyse, waarom
je deze conclusie maakt. 

Eindig je analyse met een paragraaf getiteld 'Conclusie en mitigatie'. 
Geef daarin in een paar zinnen de algemene conclusie (overzicht van verslechteringen en verbeteringen), en hoe effecten gemitigeerd kunnen worden.

"""

# --- 4. CACHED HELPER FUNCTIONS (Models & Data) ---

@st.cache_resource
def get_embedding_model():
    print("Loading embedding model...")
    return HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL, model_kwargs={'device': 'cpu'})

@st.cache_resource
def get_vector_store():
    print("Loading vector store...")
    return Chroma(persist_directory=VECTOR_STORE_DIRECTORY, embedding_function=get_embedding_model())

@st.cache_resource
def get_custom_llm():
    """LangChain wrapper voor RAG taken"""
    print("Initializing custom LLM connection...")
    return ChatOpenAI(
        model="gpt-5-mini",
        api_key=YOUR_API_KEY,
        base_url=YOUR_API_BASE_URL,
        temperature=1
    )

@st.cache_resource
def get_openai_client():
    api_key = st.session_state.get('openai_api_key')
    base_url = st.session_state.get('openai_base_url')
    
    if base_url and not base_url.endswith("/v1"):
        base_url = base_url.rstrip('/') + '/v1' 
        
    if api_key and base_url:
        return OpenAI(api_key=api_key, base_url=base_url)
    return None

@st.cache_data
def get_all_area_names():
    vector_store = get_vector_store()
    documents = vector_store.get(include=['metadatas'])
    unique_area_names = set()
    for metadata in documents['metadatas']:
        if 'area_name' in metadata:
            unique_area_names.add(metadata['area_name'])
    return sorted(list(unique_area_names))

@st.cache_data
def load_gemeenten():
    try:
        import geopandas as gpd
        if not os.path.exists(PAD_GEMEENTEN):
            return []
        gemeenten_gdf = gpd.read_file(PAD_GEMEENTEN)
        return sorted(gemeenten_gdf['Gemeentenaam'].unique().tolist())
    except Exception as e:
        print(f"Error loading gemeenten: {e}")
        return []

# --- MOCK DATA FUNCTIES (DEV MODE) ---
def get_mock_natuur_data(selected_areas):
    """Geeft instant dummy data terug voor UI testing."""
    import time
    time.sleep(0.5) # Simuleer een heel klein beetje laadtijd
    
    results = {}
    mock_json = {
        "bevindingen": [
            {"categorie": "Habitattype", "natuurtype": "H3150 Meren", "kwaliteit": "Matig", "knelpunten": "Stikstof", "oordeel": "Ja"},
            {"categorie": "Broedvogels", "natuurtype": "Roerdomp", "kwaliteit": "Slecht", "knelpunten": "Verdroging", "oordeel": "Nee, niet haalbaar"},
            {"categorie": "Habitatrichtlijnsoorten", "natuurtype": "Kamsalamander", "kwaliteit": "Onbekend", "knelpunten": "Versnippering", "oordeel": "Nee, gebrek aan gegevens"}
        ],
        "samenvatting": "Dit is een door de ontwikkelaarsmodus gegenereerde samenvatting. De kwaliteit van de meren staat onder druk door stikstof, en voor de kamsalamander ontbreken momenteel de gegevens."
    }
    
    for area in selected_areas:
        markdown_str = format_json_to_markdown(mock_json)
        results[area] = {
            'summary': markdown_str,
            'sources': ['mock_beheerplan_2024.pdf', 'mock_bijlage_kaarten.pdf'],
            'raw_data': mock_json
        }
    return results

def get_mock_impact_data():
    """Geeft instant dummy data terug voor de Omgevingsvisie UI testing."""
    return """### 1. Woningbouw
- Geplande bouw van 1500 woningen in de wijk 'Nieuw-Noord' (pag. 12).
- **Risico's:** Ruimtegebruik grenzend aan Natura 2000, extra stikstofemissie tijdens de bouwfase, en verhoogde recreatiedruk in omliggende bossen door nieuwe bewoners.

### 2. Recreatie Ontwikkeling
- Aanleg van nieuwe fietspaden en een bezoekerscentrum aan de rand van het plassengebied (pag. 45).
- **Risico's:** Verstoring van broedvogels door toename van dagjesmensen en geluid.

### 3. Mobiliteit & Infrastructuur
- Verbreding van de N-weg om nieuwe wijken te ontsluiten (pag. 50).
- **Risico's:** Extra verkeersbewegingen leiden tot hogere stikstofdepositie. Barrièrewerking (versnippering) voor grondgebonden soorten zoals dassen.

### 4. Landbouwmaatregelen
- Geen relevante ingrepen gevonden in dit document.

### 5. Bedrijvigheid
- Uitbreiding van lokaal bedrijventerrein 'De Haven' met 5 hectare (pag. 60).
- **Risico's:** Extra vrachtverkeer resulteert in stikstofuitstoot, en licht/geluidsverstoring in de nachtelijke uren voor nabijgelegen natuur."""

# --- 5. LOGIC HELPER FUNCTIONS (Matching, Processing, Conversion) ---

def load_introduction_from_docx(file_path):
    """Leest tekst uit een .docx bestand voor de inleiding."""
    if not os.path.exists(file_path):
        return None
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return "\n\n".join(full_text)
    except Exception as e:
        return f"*(Fout bij laden inleiding uit {file_path}: {e})*"

def calculate_dynamic_stopwords(all_names: list, frequency_threshold: float = 0.05):
    word_counter = Counter()
    total_docs = len(all_names)
    if total_docs == 0: return set()

    for name in all_names:
        clean = re.sub(r'[^a-z0-9\s]+', ' ', name.lower())
        words = clean.split()
        word_counter.update(set(words))
    
    dynamic_noise = set()
    cutoff_count = total_docs * frequency_threshold
    for word, count in word_counter.items():
        if count > cutoff_count and len(word) > 1:
            dynamic_noise.add(word)
    return dynamic_noise

def clean_area_name_for_matching(name: str, dynamic_stopwords: set = None) -> str:
    clean_name = name.lower()
    clean_name = re.sub(r'[^a-z0-9\s]+', ' ', clean_name)
    words = clean_name.split()
    all_stopwords = set(RUIS_WOORDEN)
    if dynamic_stopwords:
        all_stopwords.update(dynamic_stopwords)
    filtered_words = [w for w in words if w not in all_stopwords]
    clean_name = ' '.join(filtered_words)
    return re.sub(r'\s+', ' ', clean_name).strip()

def parse_json_response(response_text: str):
    try:
        cleaned_text = re.sub(r'```json\s*', '', response_text)
        cleaned_text = re.sub(r'```\s*$', '', cleaned_text)
        
        # Zoek specifiek naar het JSON object (tussen accolades) om intro/outro tekst te negeren
        match = re.search(r'(\{.*\})', cleaned_text, re.DOTALL)
        if match:
            cleaned_text = match.group(1)
        else:
            # Fallback: als regex faalt (bv. door ontbrekende sluit-accolade bij afkapping)
            start_idx = cleaned_text.find('{')
            if start_idx != -1:
                cleaned_text = cleaned_text[start_idx:]
            
        cleaned_text = cleaned_text.strip()
            
        # Reparatie voor afgekapte output (missing closing brace)
        if not cleaned_text.endswith('}'):
            cleaned_text += '}'

        # strict=False helpt bij newlines in strings die LLMs vaak genereren
        data = json.loads(cleaned_text, strict=False)
        return data
    except json.JSONDecodeError as e:
        print(f"JSON Parse Error: {e}")
        try:
            import ast
            # Fallback: ast.literal_eval is vergevingsgezinder (bv. single quotes of trailing commas)
            return ast.literal_eval(cleaned_text)
        except Exception:
            return None

def format_json_to_markdown(json_data):
    if not json_data:
        return "**Fout:** Kon geen gestructureerde data uitlezen uit het model antwoord."
    markdown_output = ""
    findings = json_data.get("bevindingen", [])
    if findings:
        df = pd.DataFrame(findings)
        expected_cols = ["categorie", "natuurtype", "kwaliteit", "knelpunten", "oordeel"]
        for col in expected_cols:
            if col not in df.columns: df[col] = ""
        rename_map = {"categorie": "Categorie", "natuurtype": "Natuurtype/Soort", "kwaliteit": "Kwaliteit", "knelpunten": "Knelpunten", "oordeel": "Eindoordeel"}
        df = df[expected_cols].rename(columns=rename_map)
        markdown_output += df.to_markdown(index=False) + "\n\n"
    else:
        markdown_output += "*Geen specifieke bevindingen gerapporteerd.*\n\n"
    summary = json_data.get("samenvatting", "")
    if summary:
        markdown_output += f"**Samenvatting:**\n{summary}"
    return markdown_output

def flatten_results_to_df(results_dict):
    rows = []
    for area, data in results_dict.items():
        raw = data.get('raw_data')
        if raw and 'bevindingen' in raw:
            for item in raw['bevindingen']:
                rows.append({
                    'Gebied': area,
                    'Categorie': item.get('categorie', 'Onbekend'),
                    'Oordeel': item.get('oordeel', 'Onbekend'),
                    'Soort': item.get('natuurtype', 'Onbekend')
                })
    return pd.DataFrame(rows)

def convert_markdown_to_docx_bytes(markdown_string: str, map_image_buffer: io.BytesIO = None) -> io.BytesIO:
    """Converteert Markdown string naar een Word document op basis van wbtemplate.docx."""
    import docx
    
    template_path = "wbtemplate.docx"
    if not os.path.exists(template_path):
        # Fallback naar leeg document als template ontbreekt
        doc = docx.Document()
    else:
        doc = docx.Document(template_path)
    
    body = doc.element.body
    
    # --- 1. Markers identificeren ---
    start_marker_para = None
    end_marker_table = None
    
    # Zoek Start Marker (Paragraaf)
    for p in doc.paragraphs:
        if "titel eerste hoofdstuk" in p.text.lower():
            start_marker_para = p
            break
            
    # Zoek End Marker (Tabel)
    for t in doc.tables:
        table_text = "".join(cell.text for row in t.rows for cell in row.cells).lower()
        if "deze tekst laten staan" in table_text and "laatste pagina berekening" in table_text:
            end_marker_table = t
            break
            
    # Als markers niet gevonden zijn, fallback naar standaard gedrag (append aan doc)
    if not start_marker_para or not end_marker_table:
        print("Waarschuwing: Markers niet gevonden in template. Gebruik fallback.")
        start_idx = len(body) - 1
        tail_elements = []
    else:
        # Indices bepalen in body
        start_idx = -1
        end_idx = -1
        for i, child in enumerate(body):
            if child is start_marker_para._element:
                start_idx = i
            if child is end_marker_table._element:
                end_idx = i
                
        if start_idx == -1 or end_idx == -1:
             start_idx = len(body) - 1
             tail_elements = []
        else:
            # 2. Tail bewaren en content tussen markers verwijderen
            tail_elements = []
            for i in range(end_idx, len(body)):
                tail_elements.append(body[i])
            
            # Verwijder alles vanaf start_idx + 1 inclusief de oude tail
            # (We voegen de tail later weer toe)
            for _ in range(len(body) - (start_idx + 1)):
                body.remove(body[start_idx + 1])

    # --- 3. Nieuwe content toevoegen ---
    parser = HtmlToDocx()
    
    # Koppel template styles aan de parser als ze bestaan
    available_styles = [s.name for s in doc.styles]
    
    # Pre-conditie: Voeg verplichte styles toe aan het document indien ze missen.
    # Dit is CRUCIAAL om te voorkomen dat htmldocx crasht en we in een retry-loop belanden.
    mandatory_styles = {
        "List Bullet": "Normal",
        "Heading 3": "Heading 2", # Wordt eigenlijk overgeslagen door prompt, maar voor veiligheid
        "Heading 4": "Normal"
    }
    for style_name, base_name in mandatory_styles.items():
        if style_name not in available_styles:
            try:
                new_style = doc.styles.add_style(style_name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                if base_name in available_styles:
                    new_style.base_style = doc.styles[base_name]
            except:
                pass

    if "Opsomming" in available_styles:
        try:
            parser.list_style = "Opsomming"
        except:
            pass

    # Voeg de kaart toe als die er is (na de start marker)
    if map_image_buffer:
        try:
            map_image_buffer.seek(0)
            doc.add_heading("Kaartoverzicht", level=1)
            doc.add_picture(map_image_buffer, width=Inches(6.0))
            doc.add_page_break()
        except Exception as e:
            print(f"Fout bij toevoegen kaart aan DOCX: {e}")

    # Converteer de rest van de markdown en voeg het toe
    html = markdown.markdown(markdown_string, extensions=['tables'])
    
    # Eén schone pass zonder retry-duplicatie
    try:
        parser.add_html_to_document(html, doc)
    except Exception as e:
        print(f"Fout bij HTML conversie: {e}")

    # --- 4. Tail terugplaatsen ---
    for elem in tail_elements:
        body.append(elem)

    # Sla op in buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def match_areas_from_csv(uploaded_file, all_available_areas: list, column_name: str = 'naam_n2k', threshold: int = 60):
    dynamic_stopwords = calculate_dynamic_stopwords(all_available_areas, frequency_threshold=0.05)
    st.session_state.dynamic_stopwords_used = sorted(list(dynamic_stopwords))

    try:
        df = pd.read_csv(uploaded_file) 
        if column_name not in df.columns:
            st.error(f"Kolom '{column_name}' niet gevonden in CSV.")
            return [], [], [] 
        
        distance_map = {}
        if 'afstand_km' in df.columns:
            for _, row in df.iterrows():
                name_key = str(row[column_name]).strip()
                try: distance_map[name_key] = float(row['afstand_km'])
                except: distance_map[name_key] = None

        csv_names = df[column_name].astype(str).str.strip().unique().tolist()
    except Exception as e:
        st.error(f"Fout bij lezen CSV: {e}")
        return [], [], []

    areas_to_analyze_indexed = set()
    successful_matches_detail = []
    debug_info = []
    
    indexed_map = {}
    for full_name in all_available_areas:
        clean_key = clean_area_name_for_matching(full_name, dynamic_stopwords)
        if clean_key: indexed_map[clean_key] = full_name
    unique_indexed_signatures = list(indexed_map.keys())

    for csv_name in csv_names:
        cleaned_csv_signature = clean_area_name_for_matching(csv_name, dynamic_stopwords)
        
        if cleaned_csv_signature in unique_indexed_signatures:
            best_match_signature = cleaned_csv_signature
            score = 100
        else:
            match_result = process.extractOne(cleaned_csv_signature, unique_indexed_signatures, scorer=fuzz.token_sort_ratio)
            best_match_signature = match_result[0] if match_result else None
            score = match_result[1] if match_result else 0
        
        original_indexed_name = indexed_map.get(best_match_signature, "Onbekend") if best_match_signature else None
        dist = distance_map.get(csv_name)

        if score >= threshold and original_indexed_name:
            areas_to_analyze_indexed.add(original_indexed_name)
            successful_matches_detail.append({'csv_name': csv_name, 'indexed_name': original_indexed_name, 'cleaned_match': f"'{cleaned_csv_signature}' == '{best_match_signature}'", 'score': score, 'distance': dist})
        else:
            debug_info.append({'csv_name': csv_name, 'best_candidate': original_indexed_name, 'cleaned_match': f"'{cleaned_csv_signature}' vs '{best_match_signature}'", 'score': score, 'distance': dist})

    return successful_matches_detail, sorted(list(areas_to_analyze_indexed)), debug_info


def analyze_local_pdf(uploaded_file, client):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name

    try:
        loader = PyPDFLoader(tmp_path)
        documents = loader.load()
        full_text = "\n\n".join([doc.page_content for doc in documents])
        
        prompt_content = IMPACT_PROMPT_FULL.format(context=full_text)
        
        response = client.chat.completions.create(
            model="gpt-5-mini", 
            messages=[
                {"role": "system", "content": "Je bent een expert in ruimtelijke ordening en ecologie."},
                {"role": "user", "content": prompt_content}
            ],
            temperature=1 
        )
        
        return response.choices[0].message.content

    except Exception as e:
        return f"Fout bij analyseren Omgevingsvisie: {e}"
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

# --- NIEUW: Functie voor het genereren van een conclusieparagraaf ---
def generate_conclusion_paragraph(topic: str, impact_analyse: str, client) -> str:
    """Genereert een enkele conclusieparagraaf voor een specifiek onderwerp op basis van impactanalyse."""
    if not client:
        return f"*(Fout: OpenAI client niet beschikbaar voor onderwerp '{topic}')*"
    
    prompt_content = CONCLUSION_PROMPT_TEMPLATE.format(
        topic=topic,
        impact_analyse=impact_analyse
    )
    
    try:
        response = client.chat.completions.create(
            model="gpt-5-mini",
            messages=[
                {"role": "system", "content": "Je bent een expert in ecologie en ruimtelijke ordening."},
                {"role": "user", "content": prompt_content}
            ],
            temperature=1 
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"*(Fout bij het genereren van conclusie voor '{topic}': {e})*"


# --- 6. RAG CHAIN FUNCTIONS ---

def create_filtered_retriever(vector_store, selected_areas):
    RETRIEVAL_COUNT = 15
    if not selected_areas:
        return vector_store.as_retriever(search_kwargs={"k": RETRIEVAL_COUNT})
    if len(selected_areas) == 1:
        chroma_filter = {"area_name": selected_areas[0]}
    else:
        chroma_filter = {"$or": [{"area_name": area} for area in selected_areas]}
    return vector_store.as_retriever(search_kwargs={"k": RETRIEVAL_COUNT, "filter": chroma_filter})

def get_rag_chain(_retriever, _llm, system_template):
    prompt = ChatPromptTemplate.from_messages([
        SystemMessagePromptTemplate.from_template(system_template),
        HumanMessagePromptTemplate.from_template("{question}") 
    ])
    return RunnableParallel({
        "context": _retriever,
        "question": RunnablePassthrough()
    }) | {
        "answer": prompt | _llm | StrOutputParser(),
        "context": lambda x: x["context"]
    }

def invoke_rag_chain(rag_chain, prompt):
    response_dict = rag_chain.invoke(prompt)
    unique_sources = set()
    for doc in response_dict['context']:
        unique_sources.add(doc.metadata.get('source', 'Onbekende bron'))
    return response_dict['answer'], sorted(list(unique_sources))

def analyze_single_area(area, vector_store, llm, system_template, concept_check_prompt, table_generation_prompt):
    """Helper function to analyze a single area. Can be run in parallel."""
    retriever = create_filtered_retriever(vector_store, [area])
    rag_chain = get_rag_chain(retriever, llm, system_template)
    try:
        concept_check_result, _ = invoke_rag_chain(rag_chain, concept_check_prompt)
        final_prompt = table_generation_prompt.format(concept_check_result=concept_check_result)
        json_response_text, unique_sources = invoke_rag_chain(rag_chain, final_prompt)
        json_data = parse_json_response(json_response_text)
        if json_data:
            formatted_markdown = format_json_to_markdown(json_data)
            return area, {'summary': formatted_markdown, 'sources': unique_sources, 'raw_data': json_data}
        else:
            return area, {'summary': f"**Fout:** Geen valide JSON.\nOutput: {json_response_text}", 'sources': unique_sources, 'raw_data': None}
    except Exception as e:
        return area, {'summary': f"Fout: {e}", 'sources': [], 'raw_data': None}

def run_batch_analysis(vector_store, llm, selected_areas, concept_check_prompt, table_generation_prompt, system_template):
    if not selected_areas: return "Geen documenten geselecteerd."
    
    results = {}
    progress_bar = st.progress(0, text="Analyse wordt voorbereid...")
    total_areas = len(selected_areas)
    
    # Use a thread pool to run analyses in parallel.
    # The number of workers can be tuned. 5 is a reasonable start.
    # The tasks are I/O bound (waiting for API), so we can use more threads than CPU cores.
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        # Create a list of future objects
        future_to_area = {
            executor.submit(
                analyze_single_area, 
                area, vector_store, llm, system_template, concept_check_prompt, table_generation_prompt
            ): area 
            for area in selected_areas
        }
        
        completed_count = 0
        # Process futures as they complete
        for future in concurrent.futures.as_completed(future_to_area):
            area = future_to_area[future]
            try:
                area_name, result_data = future.result()
                results[area_name] = result_data
            except Exception as exc:
                results[area] = {'summary': f"Fout tijdens parallelle uitvoering voor {area}: {exc}", 'sources': [], 'raw_data': None}
            
            completed_count += 1
            progress_bar.progress(completed_count / total_areas, text=f"Analyse afgerond voor: **{area}** ({completed_count}/{total_areas})")

    progress_bar.empty()
    
    # Return results in the original order of selected_areas
    ordered_results = {area: results.get(area, {'summary': 'Niet verwerkt', 'sources': [], 'raw_data': None}) for area in selected_areas}
    return ordered_results

# --- 7. MAIN APP INTERFACE ---

# --- Service Definition Class ---
# This makes the dependency explicit and clean
class AnalysisServices:
    def __init__(self, run_batch_analysis_func, analyze_local_pdf_func, get_pdf_name_func):
        self.run_batch_analysis = run_batch_analysis_func
        self.analyze_local_pdf = analyze_local_pdf_func
        self.get_pdf_name = get_pdf_name_func

st.title("🌱 Passende Beoordeling voor Omgevingsvisies en programma's")
st.markdown("Stap 1: Analyseer Natura 2000 doelen. Stap 2: Analyseer impact vanuit Omgevingsvisie.")

try:
    vector_store = get_vector_store()
    llm = get_custom_llm()
    openai_client = get_openai_client()
    all_areas = get_all_area_names()
except Exception as e:
    st.error(f"Fout bij initialisatie: {e}")
    st.stop()

# Initialize Session State
for key in ['csv_file_buffer', 'areas_to_analyze', 'locked_areas_to_analyze', 'successful_matches_detail', 'debug_info', 'analysis_results', 'dynamic_stopwords_used', 
            'natuur_analysis_md', 'impact_analysis_md', 'final_report_md', 'map_image_buffer',
            'conclusion_topics_selected', 'conclusion_results_md']: # NIEUWE KEYS
    if key not in st.session_state: st.session_state[key] = None if 'results' in key or 'buffer' in key or 'md' in key else []
if 'matching_complete' not in st.session_state: st.session_state.matching_complete = False
if 'analysis_running' not in st.session_state: st.session_state.analysis_running = False

st.sidebar.header("1. Selectie Methode")
tab1, tab2 = st.sidebar.tabs(["A. Gemeente", "B. CSV Upload"])

with tab1:
    st.subheader("Automatisch via Gemeente")
    available_gemeenten = load_gemeenten()
    gemeente_input = st.selectbox("Selecteer een Gemeente:", options=available_gemeenten, index=None, placeholder="Typ...", key='gemeente_selectbox')
    if st.button("Genereer & Match Documenten"):
        if gemeente_input:
            # Vorige kaart en selectie wissen
            st.session_state.map_image_buffer = None
            
            message, gemeente_gdf, gebieden_gdf = get_geodata_for_municipality(gemeente_input)

            if gebieden_gdf is not None and not gebieden_gdf.empty:
                st.success(f"Analyse geslaagd: {len(gebieden_gdf)} gebieden gevonden voor '{gemeente_input}'.")

                # 1. Genereer en bewaar de kaart
                with st.spinner("Kaart genereren..."):
                    map_buffer = create_map_image(gemeente_gdf, gebieden_gdf, gemeente_input)
                    st.session_state.map_image_buffer = map_buffer

                # 2. Genereer CSV voor matching-proces
                gebieden_gdf['afstand_km'] = gebieden_gdf['kortste_afstand_m'] / 1000.0
                csv_df = gebieden_gdf[['naam_n2k', 'afstand_km']]
                csv_buffer = io.BytesIO()
                csv_df.to_csv(csv_buffer, index=False)
                csv_buffer.seek(0)
                st.session_state.csv_file_buffer = csv_buffer
                
                # 3. Match gebieden
                matches, areas, debug = match_areas_from_csv(io.BytesIO(csv_buffer.getvalue()), all_areas)
                st.session_state.successful_matches_detail = matches
                st.session_state.areas_to_analyze = areas
                st.session_state.debug_info = debug
                st.session_state.locked_areas_to_analyze = list(areas)
                st.session_state.matching_complete = True
                st.rerun()
            else:
                st.warning(message)
        else:
            st.warning("Selecteer een gemeente.")

with tab2:
    st.subheader("Upload CSV")
    uploaded_file = st.file_uploader("Upload CSV met 'naam_n2k' kolom", type=['csv'])
    if uploaded_file:
        # Wis de kaart als er een nieuwe upload is
        st.session_state.map_image_buffer = None
        matches, areas, debug = match_areas_from_csv(uploaded_file, all_areas)
        st.session_state.successful_matches_detail = matches
        st.session_state.areas_to_analyze = areas
        st.session_state.debug_info = debug
        st.session_state.locked_areas_to_analyze = list(areas)
        st.session_state.matching_complete = True
        st.success(f"CSV geladen! {len(areas)} documenten.")
        st.rerun()

areas = st.session_state.areas_to_analyze
matches = st.session_state.successful_matches_detail
debug = st.session_state.debug_info

if matches or debug:
    st.sidebar.markdown("---")
    st.sidebar.subheader("Resultaten:")
    if matches:
        st.sidebar.success(f"✅ {len(matches)} matches gevonden.")
        with st.sidebar.expander(f"Details Matches ({len(matches)})"):
            for m in matches:
                dist_str = f" ({m['distance']:.1f} km)" if m.get('distance') is not None else ""
                st.write(f"**{m['csv_name']}**{dist_str}\n-> {m['indexed_name']}\n*(Op basis van: {m['cleaned_match']})*")
    
    if debug:
        st.sidebar.warning(f"❌ {len(debug)} niet gevonden.")
        with st.sidebar.expander(f"Details Mislukt ({len(debug)})"):
            for d in debug:
                dist_str = f" ({d['distance']:.1f} km)" if d.get('distance') is not None else ""
                st.write(f"**{d['csv_name']}**{dist_str}\n(Beste gok: *{d['best_candidate']}* - {d['score']}%)")

    if st.session_state.dynamic_stopwords_used:
        with st.sidebar.expander("ℹ️ Automatisch Genegeerde Ruiswoorden"):
            st.write(", ".join(st.session_state.dynamic_stopwords_used))

if areas:
    st.sidebar.markdown("---")
    st.sidebar.code(", ".join(areas), language='text')

st.info(f"Geselecteerde documenten: **{len(areas)}**")

if st.session_state.get('map_image_buffer'):
    st.subheader("Kaartoverzicht")
    st.session_state.map_image_buffer.seek(0)
    st.image(st.session_state.map_image_buffer, caption="Geselecteerde gemeente en nabijgelegen natuurgebieden.")
    st.markdown("---")

# --- DEV MODE TOGGLE ---
st.sidebar.markdown("---")
st.sidebar.subheader("🔧 Development Tools")

# Een enkele schakelaar voor dev mode. De standaardwaarde wordt bepaald door een omgevingsvariabele.
# Dit is een gebruikelijk patroon voor CI/CD of lokale tests.
# Start met: DEV_MODE=true streamlit run app7.py

# Check eerst OS environment variable, daarna secrets.toml
is_dev_env = os.getenv("DEV_MODE", "false").lower() == "true"
if not is_dev_env:
    # Fallback naar secrets.toml als env var niet is gezet
    is_dev_env = str(st.secrets.get("DEV_MODE", "false")).lower() == "true"

dev_mode_active = st.sidebar.checkbox(
    "🛠️ Ontwikkelaarsmodus (gebruik mock data)",
    value=is_dev_env
)

# --- Service Injection ---
# Op basis van de schakelaar "injecteren" we de echte of de mock services.
if dev_mode_active:
    st.sidebar.info("Dev modus is actief.")
    
    # Creëer mock services die voldoen aan de verwachte interface
    def mock_natuur_analyzer(selected_areas, **kwargs):
        st.toast("Gebruik makend van Mock Data voor Stap 1!", icon="🛠️")
        return get_mock_natuur_data(selected_areas)

    def mock_impact_analyzer(uploaded_file, **kwargs):
        return get_mock_impact_data()
        
    def get_mock_filename(uploaded_file):
        return "mock_omgevingsvisie_2024.pdf"

    services = AnalysisServices(
        run_batch_analysis_func=mock_natuur_analyzer,
        analyze_local_pdf_func=mock_impact_analyzer,
        get_pdf_name_func=get_mock_filename
    )
else:
    # Echte services
    services = AnalysisServices(
        run_batch_analysis_func=run_batch_analysis,
        analyze_local_pdf_func=analyze_local_pdf,
        get_pdf_name_func=lambda pdf: pdf.name if pdf else ""
    )

# --- EXECUTION FLOW ---

st.header("Stap 1: Natura 2000 Doelen Analyse")

if st.button("▶️ Start Stap 1 (Natuurdoelen)", disabled=not areas):
    target_areas = list(st.session_state.get('areas_to_analyze') or st.session_state.get('locked_areas_to_analyze') or [])
    if not target_areas:
        st.error("Geen gebieden geselecteerd.")
    else:
        st.session_state.analysis_running = True
        try:
            # De code hoeft niet meer te weten of het in dev-modus is.
            # Het roept simpelweg de geïnjecteerde service aan.
            results = services.run_batch_analysis(
                vector_store=vector_store, 
                llm=llm, 
                selected_areas=target_areas, 
                concept_check_prompt=CONCEPT_CHECK_PROMPT, 
                table_generation_prompt=TABLE_GENERATION_PROMPT, 
                system_template=SYSTEM_TEMPLATE
            )

            st.session_state.analysis_results = results
            
            if results:
                now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                # Inleiding laden uit docx
                intro_path = os.path.join("rapport_tekst", "INLEIDING.docx")
                intro_text = load_introduction_from_docx(intro_path)
                
                md_out = f"# Natuurdoel & Omgevings Impact Rapport\n**Datum:** {now}\n\n"
                if intro_text:
                    md_out += f"{intro_text}\n\n---\n\n"
                
                md_out += "# DEEL 1: Natura 2000 Analyse\n\n"
                for area, data in results.items():
                    md_out += f"## {area}\n\n{data['summary']}\n\n"
                    if data['sources']: md_out += "**Bronnen:**\n" + "\n".join([f"- {s}" for s in data['sources']]) + "\n"
                    md_out += "\n---\n\n"
                
                st.session_state.natuur_analysis_md = md_out
                st.session_state.final_report_md = md_out 
                st.success("Stap 1 Voltooid!")
                st.rerun()
        except Exception as e:
            st.error(f"Fout: {e}")
        finally:
            st.session_state.analysis_running = False

if st.session_state.analysis_results:
    with st.expander("Bekijk resultaten Stap 1", expanded=False):
        st.markdown(st.session_state.natuur_analysis_md)

    st.markdown("---")
    st.header("Stap 2: Omgevingsvisie Impact Analyse (Optioneel)")
    st.markdown("Upload de Omgevingsvisie (PDF) om te controleren op ingrepen die de geselecteerde natuurgebieden kunnen raken.")

    omgevings_pdf = st.file_uploader("Upload Omgevingsvisie PDF", type="pdf", key="omgevingsvisie_uploader")

    # Als DEV MODE aanstaat, sta toe dat we Stap 2 draaien zónder een PDF te hoeven uploaden
    can_run_step_2 = (omgevings_pdf is not None) or dev_mode_active

    if can_run_step_2 and st.button("▶️ Start Stap 2 (Impact Analyse)"):
        if not openai_client and not dev_mode_active:
            st.error("Kon OpenAI client niet laden. Controleer je secrets.")
        else:
            with st.spinner("Bezig met analyseren van volledige Omgevingsvisie..."):
                # Roep de juiste service aan (echt of mock)
                impact_result = services.analyze_local_pdf(
                    uploaded_file=omgevings_pdf, 
                    client=openai_client
                )
                bestandsnaam = services.get_pdf_name(omgevings_pdf)

                impact_md = "\n# DEEL 2: Ingreep-effect relaties\n\n"
                impact_md += f"**Geanalyseerd bestand:** {bestandsnaam}\n\n"
                impact_md += impact_result

                st.session_state.impact_analysis_md = impact_md
                
                st.session_state.final_report_md = st.session_state.natuur_analysis_md + "\n\n---\n\n" + impact_md
                st.success("Stap 2 Voltooid! Rapport bijgewerkt.")
                st.rerun()

    if st.session_state.impact_analysis_md:
        with st.expander("Bekijk resultaten Stap 2", expanded=True):
            st.markdown(st.session_state.impact_analysis_md)

        # --- NIEUW: STAP 3 ---
        st.markdown("---")
        st.header("Stap 3: Conclusies Genereren (Optioneel)")
        st.markdown("Selecteer de onderwerpen waarover u een concluderende paragraaf wilt genereren op basis van de impactanalyse (Stap 2).")

        # NIEUW: Lijst met onderwerpen voor de conclusie
        CONCLUSION_TOPICS = [
            "Stikstofdepositie", "Recreatiedruk en toename verstoring", "Behoud en versterken van natuur buiten Natura 2000 gebieden",
            "Algemene Samenvatting & Aanbevelingen"
        ]

        selected_topics = st.multiselect(
            "Kies onderwerpen voor de conclusie:",
            options=CONCLUSION_TOPICS,
            key='conclusion_topics_selected'
        )

        if st.button("▶️ Start Stap 3 (Genereer Conclusies)", disabled=not selected_topics):
            if not openai_client and not dev_mode_active:
                st.error("Kon OpenAI client niet laden. Controleer je secrets.")
            else:
                conclusion_results = {}
                total_topics = len(selected_topics)
                progress_bar = st.progress(0, text="Voorbereiden van conclusies...")

                with st.spinner("Bezig met genereren van conclusies..."):
                    for i, topic in enumerate(selected_topics):
                        progress_bar.progress((i + 1) / total_topics, text=f"Bezig met: **{topic}** ({i+1}/{total_topics})")
                        
                        # In dev mode, gebruik een mock antwoord
                        if dev_mode_active:
                            import time
                            time.sleep(0.5)
                            conclusion_text = f"Dit is een mock-conclusie voor het onderwerp **{topic}**. De impactanalyse toont aan dat de geplande ontwikkelingen significante risico's met zich meebrengen."
                        else:
                            conclusion_text = generate_conclusion_paragraph(
                                topic=topic,
                                impact_analyse=st.session_state.impact_analysis_md,
                                client=openai_client
                            )
                        conclusion_results[topic] = conclusion_text

                progress_bar.empty()
                
                # Formatteer de resultaten naar Markdown
                conclusion_md = "\n# DEEL 3: Conclusies\n\n"
                for topic, text in conclusion_results.items():
                    conclusion_md += f"### Conclusie: {topic}\n"
                    conclusion_md += f"{text}\n\n"
                
                st.session_state.conclusion_results_md = conclusion_md
                
                # Werk het volledige rapport bij
                st.session_state.final_report_md = (st.session_state.natuur_analysis_md + "\n\n---\n\n" + st.session_state.impact_analysis_md + "\n\n---\n\n" + conclusion_md)
                st.success("Stap 3 Voltooid! Rapport bijgewerkt met conclusies.")
                st.rerun()

    if st.session_state.get('conclusion_results_md'):
        with st.expander("Bekijk resultaten Stap 3", expanded=True):
            st.markdown(st.session_state.conclusion_results_md)

# --- DOWNLOADS & STATS ---

if st.session_state.final_report_md:
    st.markdown("---")
    st.header("📥 Download Eindrapport")
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    c1, c2 = st.columns(2)
    c1.download_button("⬇️ Download Volledig Rapport (.md)", st.session_state.final_report_md, f"rapport_compleet_{timestamp}.md")
    
    docx_data = convert_markdown_to_docx_bytes(st.session_state.final_report_md, st.session_state.get('map_image_buffer'))
    c2.download_button("⬇️ Download Volledig Rapport (.docx)", docx_data, f"rapport_compleet_{timestamp}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if st.session_state.analysis_results:
    df_stats = flatten_results_to_df(st.session_state.analysis_results)
    if not df_stats.empty:
        st.markdown("---")
        st.header("📊 Kwantitatieve Analyse (Natura 2000)")
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Aantal beoordelingen per Categorie")
            st.bar_chart(df_stats['Categorie'].value_counts())
        with col2:
            st.subheader("Verdeling van Oordelen")
            st.bar_chart(pd.crosstab(df_stats['Categorie'], df_stats['Oordeel']))