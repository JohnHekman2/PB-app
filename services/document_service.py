import os
import io

def load_introduction_from_docx(file_path):
    """Leest tekst uit een .docx bestand voor de inleiding."""
    if not os.path.exists(file_path):
        return None
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        return "\n\n".join(full_text)
    except Exception as e:
        return f"*(Fout bij laden inleiding uit {file_path}: {e})*"

def convert_markdown_to_docx_bytes(markdown_string: str, map_image_buffer: io.BytesIO = None) -> io.BytesIO:
    """Converteert Markdown string naar een Word document op basis van wbtemplate.docx."""
    import docx
    from htmldocx import HtmlToDocx
    from docx.shared import Inches
    import markdown
    
    template_path = "wbtemplate.docx"
    if not os.path.exists(template_path):
        # Fallback naar leeg document als template ontbreekt
        doc = docx.Document()
    else:
        doc = docx.Document(template_path)
    
    body = doc.element.body
    
    # --- 1. Markers identificeren ---
    start_marker_para = None
    end_marker_table = None
    
    # Zoek Start Marker (Paragraaf)
    for p in doc.paragraphs:
        if "titel eerste hoofdstuk" in p.text.lower():
            start_marker_para = p
            break
            
    # Zoek End Marker (Tabel)
    for t in doc.tables:
        table_text = "".join(cell.text for row in t.rows for cell in row.cells).lower()
        if "deze tekst laten staan" in table_text and "laatste pagina berekening" in table_text:
            end_marker_table = t
            break
            
    # Als markers niet gevonden zijn, fallback naar standaard gedrag (append aan doc)
    if not start_marker_para or not end_marker_table:
        print("Waarschuwing: Markers niet gevonden in template. Gebruik fallback.")
        start_idx = len(body) - 1
        tail_elements = []
    else:
        # Indices bepalen in body
        start_idx = -1
        end_idx = -1
        for i, child in enumerate(body):
            if child is start_marker_para._element:
                start_idx = i
            if child is end_marker_table._element:
                end_idx = i
                
        if start_idx == -1 or end_idx == -1:
             start_idx = len(body) - 1
             tail_elements = []
        else:
            # 2. Tail bewaren en content tussen markers verwijderen
            tail_elements = []
            for i in range(end_idx, len(body)):
                tail_elements.append(body[i])
            
            # Verwijder alles vanaf start_idx + 1 inclusief de oude tail
            # (We voegen de tail later weer toe)
            for _ in range(len(body) - (start_idx + 1)):
                body.remove(body[start_idx + 1])

    # --- 3. Nieuwe content toevoegen ---
    parser = HtmlToDocx()
    
    # Koppel template styles aan de parser als ze bestaan
    available_styles = [s.name for s in doc.styles]
    
    # Pre-conditie: Voeg verplichte styles toe aan het document indien ze missen.
    # Dit is CRUCIAAL om te voorkomen dat htmldocx crasht en we in een retry-loop belanden.
    mandatory_styles = {
        "List Bullet": "Normal",
        "Heading 3": "Heading 2", # Wordt eigenlijk overgeslagen door prompt, maar voor veiligheid
        "Heading 4": "Normal"
    }
    for style_name, base_name in mandatory_styles.items():
        if style_name not in available_styles:
            try:
                new_style = doc.styles.add_style(style_name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
                if base_name in available_styles:
                    new_style.base_style = doc.styles[base_name]
            except:
                pass

    if "Opsomming" in available_styles:
        try:
            parser.list_style = "Opsomming"
        except:
            pass

    # Voeg de kaart toe als die er is (na de start marker)
    if map_image_buffer:
        try:
            map_image_buffer.seek(0)
            doc.add_heading("Kaartoverzicht", level=1)
            doc.add_picture(map_image_buffer, width=Inches(6.0))
            doc.add_page_break()
        except Exception as e:
            print(f"Fout bij toevoegen kaart aan DOCX: {e}")

    # Converteer de rest van de markdown en voeg het toe
    html = markdown.markdown(markdown_string, extensions=['tables'])
    
    # Eén schone pass zonder retry-duplicatie
    try:
        parser.add_html_to_document(html, doc)
    except Exception as e:
        print(f"Fout bij HTML conversie: {e}")

    # --- 4. Tail terugplaatsen ---
    for elem in tail_elements:
        body.append(elem)

    # Sla op in buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
